"""RAG-1.1: multi-format ingestion.

Production target is Unstructured.io (layout-aware: tables, numbered
clauses survive extraction — see requirements/03-rag-pipeline-requirements.md
and requirements-full.txt). The parsers below cover the same format matrix
(Markdown/plain text, HTML circulars, PDF, DOCX) using lightweight libraries
so the MVP ingests real files without the heavier dependency tree; each
parser is isolated behind `parse_document()` so swapping in Unstructured.io
later only touches this module.
"""

from dataclasses import dataclass
from pathlib import Path

from rag.exceptions import UnparseableDocumentError

_MIN_TEXT_LENGTH = 20


@dataclass
class ParsedDocument:
    raw_text: str
    front_matter: dict
    source_uri: str


def parse_document(path: str | Path) -> ParsedDocument:
    path = Path(path)
    if not path.exists():
        raise UnparseableDocumentError(str(path), "file not found")

    suffix = path.suffix.lower()
    try:
        if suffix in {".md", ".txt"}:
            return _parse_markdown(path)
        if suffix == ".html" or suffix == ".htm":
            return _parse_html(path)
        if suffix == ".pdf":
            return _parse_pdf(path)
        if suffix == ".docx":
            return _parse_docx(path)
    except UnparseableDocumentError:
        raise
    except Exception as exc:  # noqa: BLE001 — RAG-7.1: any parser failure is quarantined, not fatal
        raise UnparseableDocumentError(str(path), f"{type(exc).__name__}: {exc}") from exc

    raise UnparseableDocumentError(str(path), f"unsupported extension '{suffix}'")


def _parse_markdown(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8")
    front_matter, body = _split_front_matter(text)
    _guard_min_length(body, path)
    return ParsedDocument(raw_text=body, front_matter=front_matter, source_uri=str(path))


def _split_front_matter(text: str) -> tuple[dict, str]:
    """Sample corpus documents carry a YAML front-matter block (RAG-1.3
    metadata) delimited by `---` lines, avoiding a bespoke sidecar format.
    """
    import yaml

    if not text.startswith("---"):
        return {}, text
    parts = text.split("---", 2)
    if len(parts) < 3:
        return {}, text
    _, fm_text, body = parts
    front_matter = yaml.safe_load(fm_text) or {}
    return front_matter, body.strip()


def _parse_html(path: Path) -> ParsedDocument:
    from bs4 import BeautifulSoup

    html = path.read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "html.parser")
    text = soup.get_text(separator="\n")
    _guard_min_length(text, path)
    return ParsedDocument(raw_text=text, front_matter={}, source_uri=str(path))


def _parse_pdf(path: Path) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages)
    if len(text.strip()) < _MIN_TEXT_LENGTH:
        # RAG-1.1 OCR fallback hook: a scanned PDF yields ~no extractable
        # text. Production wires Unstructured.io's `hi_res`/OCR strategy
        # here; flagged rather than silently returning near-empty content.
        raise UnparseableDocumentError(
            str(path),
            "near-empty text extraction — looks scanned; OCR fallback not "
            "enabled in this MVP build (see requirements-full.txt: unstructured[pdf])",
        )
    return ParsedDocument(raw_text=text, front_matter={}, source_uri=str(path))


def _parse_docx(path: Path) -> ParsedDocument:
    import docx

    document = docx.Document(str(path))
    text = "\n".join(p.text for p in document.paragraphs)
    _guard_min_length(text, path)
    return ParsedDocument(raw_text=text, front_matter={}, source_uri=str(path))


def _guard_min_length(text: str, path: Path) -> None:
    if len(text.strip()) < _MIN_TEXT_LENGTH:
        raise UnparseableDocumentError(str(path), "extracted text is empty or too short")
